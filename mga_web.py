# ======================================================
# MGA IA WEB – SISTEMA COMPLETO (ZIP TEMPORAL)
# ======================================================
import os, json, re, io, zipfile
import pandas as pd
from fastapi import FastAPI, Form
from fastapi.responses import HTMLResponse, StreamingResponse
from docx import Document

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter


# ======================================================
# 🔐 CONFIGURACIÓN
# ======================================================
MODEL = "gpt-4.1"
TEMPERATURE = 0.2

BASE = "data"
PDF_MGA = f"{BASE}/pdf_mga_ejemplos"
DOC_BASE = f"{BASE}/documento_tecnico_base"
PDD = f"{BASE}/plan_desarrollo"

llm = ChatOpenAI(
    model=MODEL,
    temperature=TEMPERATURE,
)

# ======================================================
# 📚 CARGA DEL CORPUS (RAG)
# ======================================================
def cargar_corpus():
    documentos = []
    for carpeta in [PDF_MGA, DOC_BASE, PDD]:
        if not os.path.exists(carpeta):
            continue
        for archivo in os.listdir(carpeta):
            if archivo.endswith(".pdf"):
                documentos.extend(
                    PyPDFLoader(os.path.join(carpeta, archivo)).load()
                )

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1200,
        chunk_overlap=200
    )

    chunks = splitter.split_documents(documentos)
    return FAISS.from_documents(chunks, OpenAIEmbeddings())

db = cargar_corpus()

# ======================================================
# 🧠 IA – FORMULADOR MGA (ROBUSTO)
# ======================================================
def extraer_json_seguro(texto: str) -> dict:
    """
    Extrae el primer JSON válido aunque el modelo
    agregue texto antes o después.
    """
    match = re.search(r"\{[\s\S]*\}", texto)
    if not match:
        raise ValueError("No se encontró JSON en la respuesta")
    return json.loads(match.group())

def consultar_mga(descripcion: str) -> dict:
    contexto = db.similarity_search(descripcion, k=12)
    contexto_txt = "\n\n".join(c.page_content for c in contexto)

    prompt = f"""
Eres un FORMULADOR EXPERTO en Metodología General Ajustada (MGA) – Colombia.

USA EXCLUSIVAMENTE:
- Proyectos MGA en PDF
- Documento Técnico Base
- Plan de Desarrollo Departamental Cauca 2024–2027

REQUISITOS OBLIGATORIOS:
- Presupuesto ALTO y realista (incluye personal, capacitadores, operación)
- Costos en COP
- Lenguaje institucional MGA
- Estructura MGA completa
- NO inventes normas
- NO expliques nada
- NO uses markdown
- NO agregues texto fuera del JSON

RESPONDE ÚNICAMENTE CON JSON VÁLIDO:

{{
  "documento_tecnico": "texto completo",
  "cadena_valor": [
    {{"Actividad": "", "Producto": "", "Costo_COP": ""}}
  ],
  "concepto_sectorial": [
    {{"Sector": "", "Alineacion_PND": ""}}
  ],
  "mga_txt": "formulación MGA completa en texto plano"
}}

CONTEXTO:
{contexto_txt}

DESCRIPCIÓN DEL PROYECTO:
{descripcion}
"""

    respuesta = llm.invoke(prompt).content
    return extraer_json_seguro(respuesta)

# ======================================================
# 📄 GENERADORES EN MEMORIA
# ======================================================
def generar_docx(texto: str) -> bytes:
    doc = Document()
    doc.add_heading("DOCUMENTO TÉCNICO MGA", 0)
    for bloque in texto.split("\n\n"):
        doc.add_paragraph(bloque)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def generar_excel(data: list) -> bytes:
    df = pd.DataFrame(data)
    buffer = io.BytesIO()
    df.to_excel(buffer, index=False)
    buffer.seek(0)
    return buffer.read()

def generar_txt(texto: str) -> bytes:
    return texto.encode("utf-8")

# ======================================================
# 📦 ZIP TEMPORAL
# ======================================================
def generar_zip(data: dict) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        zipf.writestr("Documento_Tecnico_MGA.docx", generar_docx(data["documento_tecnico"]))
        zipf.writestr("CADENA_VALOR.xlsx", generar_excel(data["cadena_valor"]))
        zipf.writestr("CONCEPTO_SECTORIAL.xlsx", generar_excel(data["concepto_sectorial"]))
        zipf.writestr("Proyecto_MGA.txt", generar_txt(data["mga_txt"]))
    buffer.seek(0)
    return buffer.read()

# ======================================================
# 🌐 WEB APP
# ======================================================
app = FastAPI(title="MGA IA Web", version="1.0")

@app.get("/", response_class=HTMLResponse)
def home():
    return """
<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>MGA IA – Fundación Almagua</title>
<meta name="viewport" content="width=device-width, initial-scale=1.0">

<script src="https://cdn.jsdelivr.net/npm/gsap@3.12.5/dist/gsap.min.js"></script>

<style>
:root{
  --bg-main:#1f2a2e;
  --bg-soft:#2e3c41;
  --bg-deep:#141c1f;
  --gold:#d4af37;
  --gold-light:#f6e27a;
}

*{box-sizing:border-box}

body{
  margin:0;
  font-family:system-ui,sans-serif;
  background:var(--bg-main);
  color:white;
  overflow-x:hidden;
}

/* ================= HEADER ================= */
header{
  position:fixed;
  top:0;
  width:100%;
  padding:18px 0 14px;
  text-align:center;
  z-index:1000;
  background:linear-gradient(
    180deg,
    rgba(20,28,31,.95),
    rgba(20,28,31,.5),
    transparent
  );
  backdrop-filter:blur(6px);
}

.logo{
  max-width:150px;
  display:inline-block;
  filter:drop-shadow(0 6px 18px rgba(0,0,0,.45));
}

/* ================= SECTION ================= */
.section{
  position:relative;
  min-height:100vh;
  padding-top:140px;
  display:flex;
  align-items:center;
  justify-content:center;
  text-align:center;
  background:
    radial-gradient(circle at center,var(--bg-soft),var(--bg-main),var(--bg-deep));
  overflow:hidden;
}

/* ===== MAIN TEXT ===== */
.main-text{
  font-size:clamp(2.4rem,6vw,3.8rem);
  font-weight:800;
  line-height:1.05;
  background:linear-gradient(45deg,var(--gold-light),var(--gold));
  -webkit-background-clip:text;
  -webkit-text-fill-color:transparent;
  transition:opacity .3s ease;
  z-index:2;
}

/* ===== REVEAL MASK ===== */
.reveal-layer{
  --x:50%;
  --y:50%;
  --r:0px;

  position:absolute;
  inset:0;
  display:flex;
  align-items:center;
  justify-content:center;
  text-align:center;

  background:linear-gradient(180deg,var(--gold-light),#8fb7c8);
  color:var(--bg-deep);

  -webkit-mask-image:
    radial-gradient(circle at var(--x) var(--y),
      black var(--r), transparent 0);
  mask-image:
    radial-gradient(circle at var(--x) var(--y),
      black var(--r), transparent 0);

  pointer-events:none;
}

.reveal-layer p{
  font-size:clamp(2.4rem,6vw,3.6rem);
  font-weight:800;
  line-height:1.05;
  margin:0;
  max-width:900px;
}

/* ================= CTA ================= */
.cta{
  position:fixed;
  bottom:30px;
  right:30px;
  background:linear-gradient(45deg,var(--gold-light),var(--gold));
  color:#141c1f;
  border:none;
  padding:16px 26px;
  border-radius:40px;
  font-weight:bold;
  cursor:pointer;
  z-index:900;
}

/* ================= POPUP ================= */
.overlay{
  position:fixed;
  inset:0;
  background:rgba(0,0,0,.6);
  backdrop-filter:blur(8px);
  display:none;
  align-items:center;
  justify-content:center;
  z-index:2000;
}

.popup{
  background:#fff;
  color:#1f2a2e;
  width:90%;
  max-width:520px;
  border-radius:22px;
  padding:26px;
  position:relative;
  box-shadow:0 40px 80px rgba(0,0,0,.45);
}

.popup h3{
  margin:0 0 8px;
}

.popup p{
  margin:0 0 14px;
  font-size:.95rem;
}

.popup textarea{
  width:100%;
  height:140px;
  padding:14px;
  border-radius:14px;
  border:1px solid #ccc;
  resize:none;
}

.popup button{
  width:100%;
  margin-top:14px;
  background:linear-gradient(45deg,var(--gold-light),var(--gold));
  border:none;
  padding:14px;
  border-radius:30px;
  font-weight:bold;
  cursor:pointer;
}

.close{
  position:absolute;
  top:14px;
  right:18px;
  cursor:pointer;
  font-weight:bold;
  opacity:.6;
}
</style>
</head>

<body>

<header>
  <img src="assets/logo.png" class="logo" alt="Fundación Almagua">
</header>

<section class="section">
  <div class="main-text">
    Las comunidades tienen ideas.<br>
    El reto siempre fue estructurarlas.
  </div>

  <div class="reveal-layer">
    <p>
      IA MGA para transformar<br>
      ideas en proyectos viables.
    </p>
  </div>
</section>

<button class="cta" id="openPopup">Generar Proyecto MGA</button>

<!-- POPUP -->
<div class="overlay" id="overlay">
  <div class="popup">
    <div class="close" id="closePopup">✕</div>
    <h3>Generador MGA con IA</h3>
    <p>Describe la idea. La estructura la construimos contigo.</p>
    <textarea placeholder="Ej: Proyecto comunitario de agua y empleo rural..."></textarea>
    <button>Generar Proyecto</button>
  </div>
</div>

<script>
/* ===== MASK EFFECT ===== */
const section = document.querySelector(".section");
const reveal = section.querySelector(".reveal-layer");
const mainText = section.querySelector(".main-text");

section.addEventListener("mousemove", e=>{
  const r = section.getBoundingClientRect();
  gsap.to(reveal,{
    "--x":(e.clientX-r.left)+"px",
    "--y":(e.clientY-r.top)+"px",
    "--r":"220px",
    duration:.35,
    ease:"power3.out"
  });
  mainText.style.opacity=.15;
});

section.addEventListener("mouseleave",()=>{
  gsap.to(reveal,{
    "--r":"0px",
    duration:.4,
    ease:"power3.inOut"
  });
  mainText.style.opacity=1;
});

/* ===== POPUP ===== */
const overlay=document.getElementById("overlay");
document.getElementById("openPopup").onclick=()=>{
  overlay.style.display="flex";
  gsap.fromTo(".popup",{scale:.9,opacity:0},{scale:1,opacity:1,duration:.4});
};

document.getElementById("closePopup").onclick=()=>{
  gsap.to(".popup",{scale:.9,opacity:0,duration:.3,onComplete:()=>{
    overlay.style.display="none";
  }});
};
</script>

</body>
</html>
"""




@app.post("/generar")
def generar(descripcion: str = Form(...)):
    data = consultar_mga(descripcion)
    zip_bytes = generar_zip(data)
    return StreamingResponse(
        io.BytesIO(zip_bytes),
        media_type="application/zip",
        headers={
            "Content-Disposition": "attachment; filename=Proyecto_MGA_Completo.zip"
        }
    )
