# ======================================================
# MGA IA WEB – SISTEMA COMPLETO (FORMULADOR SERIO)
# Fundación Almagua
# ======================================================

import os, json, re, io, zipfile
import pandas as pd
from fastapi import FastAPI, Form
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from docx import Document
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ======================================================
# 🌐 APP
# ======================================================
app = FastAPI(title="MGA IA – Fundación Almagua", version="1.0")
app.mount("/assets", StaticFiles(directory="assets"), name="assets")

# ======================================================
# 🔐 CONFIGURACIÓN
# ======================================================
MODEL = "gpt-4.1"
TEMPERATURE = 0.2
BASE = "data"

PDF_MGA = f"{BASE}/pdf_mga_ejemplos"
DOC_BASE = f"{BASE}/documento_tecnico_base"
PDD = f"{BASE}/plan_desarrollo"

llm = ChatOpenAI(model=MODEL, temperature=TEMPERATURE)

# ======================================================
# 📚 CORPUS RAG (PDFs)
# ======================================================
def cargar_corpus():
    documentos = []
    for carpeta in [PDF_MGA, DOC_BASE, PDD]:
        if not os.path.exists(carpeta):
            continue
        for archivo in os.listdir(carpeta):
            path = os.path.join(carpeta, archivo)
            if archivo.lower().endswith(".pdf"):
                documentos.extend(PyPDFLoader(path).load())
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)
    chunks = splitter.split_documents(documentos)
    return FAISS.from_documents(chunks, OpenAIEmbeddings())

db = cargar_corpus()

# ======================================================
# 🧰 CACHE DE INFORMACIÓN
# ======================================================
def generar_cache_completo():
    """
    Genera un cache de información clave que GPT usará para completar
    el MGA sin inventar datos irrelevantes.
    Incluye PDFs, DOCX y CSVs relevantes (gestion_social.csv, mujer.csv, etc.)
    """
    cache = {}

    # Cargar CSVs clave
    for csv_file in ["gestion_social.csv", "mujer.csv"]:
        path_csv = os.path.join(BASE, csv_file)
        if os.path.exists(path_csv):
            df = pd.read_csv(path_csv)
            cache[csv_file] = df.to_dict(orient="records")

    # PDFs y DOCX de referencia
    for carpeta in [PDF_MGA, DOC_BASE, PDD]:
        if not os.path.exists(carpeta):
            continue
        for archivo in os.listdir(carpeta):
            path = os.path.join(carpeta, archivo)
            if archivo.lower().endswith(".pdf"):
                loader = PyPDFLoader(path)
                docs = loader.load()
                cache[archivo] = [d.page_content for d in docs]
    return cache

cache_proyecto = generar_cache_completo()

# ======================================================
# 🧠 FUNCIONES AUXILIARES
# ======================================================
def extraer_json_seguro(respuesta: str) -> dict:
    """
    Extrae un JSON de manera segura del string generado por el modelo.
    """
    try:
        return json.loads(respuesta)
    except json.JSONDecodeError:
        cleaned = re.search(r"\{.*\}", respuesta, re.DOTALL)
        if cleaned:
            return json.loads(cleaned.group())
        else:
            return {}

# ======================================================
# 🧠 IA MGA – JSON SEGURO AVANZADO
# ======================================================
def consultar_mga(descripcion: str) -> dict:
    """
    Genera el MGA completo usando:
    - Cache de documentos clave
    - Información de CSV para presupuestos y costos
    - Ejemplo MGA con todos los campos oficiales
    """
    contexto = db.similarity_search(descripcion, k=10)
    contexto_txt = "\n\n".join(c.page_content for c in contexto)
    if len(contexto_txt) > 100_000:
        contexto_txt = contexto_txt[-100_000:]

    prompt = f"""
Eres un FORMULADOR EXPERTO MGA – Colombia.
Tu tarea es generar un MGA COMPLETO, siguiendo estrictamente la estructura oficial y los ejemplos.

SECCIONES A GENERAR (obligatorio completar todas):
1. Datos básicos del proyecto
2. Contribución al Plan Nacional de Desarrollo
3. Plan de Desarrollo Departamental o Sectorial
4. Plan de Desarrollo Distrital o Municipal
5. Identificación y descripción del problema
6. Identificación y análisis de participantes
7. Población afectada y objetivo
8. Objetivos generales y específicos con indicadores
9. Alternativas de solución
10. Estudio de necesidades
11. Análisis técnico de la alternativa
12. Localización de la alternativa
13. Cadena de valor
14. Análisis de riesgos
15. Flujo económico
16. Indicadores y decisión
17. Esquema financiero y clasificación presupuestal
18. Resumen del proyecto

REGLAS IMPORTANTES:
- No repitas nombres de personas ni información innecesaria de los documentos.
- Usa la información de los CSVs clave (gestion_social.csv, mujer.csv) para calcular presupuestos, costos, ingresos y beneficios.
- Completa todos los campos oficiales con información coherente y basada en los documentos de referencia y cache.
- Responde SOLO en JSON válido con este formato:

{{
    "documento_tecnico": "",
    "cadena_valor": [{{}}],
    "concepto_sectorial": [{{}}],
    "mga_txt": ""
}}

FUENTES DE INFORMACIÓN:
- Cache de proyecto: {json.dumps(cache_proyecto)[:3000]}  # solo preview de referencia
- Contexto similar: {contexto_txt}

DESCRIPCIÓN DEL PROYECTO:
{descripcion}
"""
    respuesta = llm.invoke(prompt).content
    return extraer_json_seguro(respuesta)

# ======================================================
# 📄 GENERADORES DE ARCHIVOS COMPLETOS
# ======================================================
def generar_docx(texto):
    doc = Document()
    doc.add_heading("DOCUMENTO TÉCNICO MGA", 0)
    if isinstance(texto, dict):
        for seccion, contenido in texto.items():
            doc.add_heading(seccion, level=1)
            if isinstance(contenido, (dict, list)):
                doc.add_paragraph(json.dumps(contenido, ensure_ascii=False, indent=2))
            else:
                doc.add_paragraph(str(contenido))
    elif isinstance(texto, list):
        doc.add_paragraph(json.dumps(texto, ensure_ascii=False, indent=2))
    else:
        doc.add_paragraph(str(texto))
    b = io.BytesIO()
    doc.save(b)
    b.seek(0)
    return b.read()

def generar_csv(data):
    if not data:
        return b""
    df = pd.DataFrame(data)
    b = io.BytesIO()
    df.to_csv(b, index=False, encoding='utf-8-sig')
    b.seek(0)
    return b.read()

def generar_zip_completo(data):
    b = io.BytesIO()
    with zipfile.ZipFile(b, "w", zipfile.ZIP_DEFLATED) as z:
        doc_tecnico = data.get("documento_tecnico", "")
        # si es dict, extrae la sección 'mga_txt'
        if isinstance(doc_tecnico, dict):
            doc_tecnico = doc_tecnico.get("mga_txt", "")
        z.writestr("Documento_Tecnico_MGA.docx", generar_docx(doc_tecnico))
        z.writestr("CADENA_VALOR.csv", generar_csv(data.get("cadena_valor", [])))
        z.writestr("CONCEPTO_SECTORIAL.csv", generar_csv(data.get("concepto_sectorial", [])))
        z.writestr("Proyecto_MGA.txt", str(data.get("mga_txt", "")))
    b.seek(0)
    return b.read()

# ======================================================
# 🌐 INTERFAZ WEB
# ======================================================

@app.get("/", response_class=HTMLResponse)
def home():
    return """
<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>MGA IA – Fundación Almagua</title>
<meta name="viewport" content="width=device-width, initial-scale=1">
<script src="https://cdn.jsdelivr.net/npm/gsap@3.12.5/dist/gsap.min.js"></script>
<style>
:root{
--bg:#1f2a2e;
--deep:#141c1f;
--gold:#d4af37;
--gold2:#f6e27a;
}
*{box-sizing:border-box}
body{
margin:0;
background:radial-gradient(circle,var(--bg),var(--deep));
color:white;
font-family:system-ui;
}
header{
position:fixed;
top:0;
width:100%;
padding:16px;
text-align:center;
background:rgba(20,28,31,.85);
backdrop-filter:blur(8px);
z-index:1000;
}
.logo{max-width:160px}
.section{
min-height:100vh;
display:flex;
align-items:center;
justify-content:center;
text-align:center;
padding-top:140px;
position:relative;
}
.main-text{
font-size:clamp(2.6rem,6vw,3.8rem);
font-weight:800;
background:linear-gradient(45deg,var(--gold2),var(--gold));
-webkit-background-clip:text;
-webkit-text-fill-color:transparent;
transition:.3s;
}
.reveal{
position:absolute;
inset:0;
display:flex;
align-items:center;
justify-content:center;
font-size:clamp(2.6rem,6vw,3.6rem);
font-weight:800;
color:#141c1f;
background:linear-gradient(180deg,var(--gold2),#8fb7c8);
mask-image:radial-gradient(circle at var(--x,50%) var(--y,50%),black var(--r,0),transparent 0);
-webkit-mask-image:radial-gradient(circle at var(--x,50%) var(--y,50%),black var(--r,0),transparent 0);
pointer-events:none;
}
.cta{
position:fixed;
bottom:30px;
right:30px;
padding:16px 28px;
border-radius:40px;
border:none;
background:linear-gradient(45deg,var(--gold2),var(--gold));
font-weight:700;
cursor:pointer;
}
.overlay{
position:fixed;
inset:0;
display:none;
align-items:center;
justify-content:center;
background:rgba(0,0,0,.65);
backdrop-filter:blur(8px);
z-index:2000;
}
.popup{
background:white;
color:#1f2a2e;
padding:32px;
border-radius:22px;
width:92%;
max-width:520px;
display:grid;
grid-template-rows:auto auto 1fr auto;
gap:18px;
}
.popup-header{
display:flex;
justify-content:space-between;
align-items:center;
}
.popup-header h3{margin:0}
.popup-header button{
background:none;
border:none;
font-size:1.4rem;
cursor:pointer;
opacity:.6;
}
.popup p{margin:0;font-size:.95rem;opacity:.85}
.popup form{
display:grid;
gap:16px;
}
textarea{
width:100%;
min-height:140px;
padding:16px;
border-radius:14px;
border:1px solid #ccc;
resize:none;
font-family:system-ui;
}
textarea:focus{
outline:none;
border-color:var(--gold);
}
button.submit{
padding:14px;
border-radius:30px;
border:none;
background:linear-gradient(45deg,var(--gold2),var(--gold));
font-weight:700;
cursor:pointer;
}
</style>
</head>
<body>
<header>
<img src="/assets/logo.png" class="logo" alt="Fundación Almagua">
</header>
<section class="section" id="hero">
<div class="main-text">Las comunidades tienen ideas.<br>La IA las convierte en MGA.</div>
<div class="reveal">IA territorial para proyectos reales.</div>
</section>
<button class="cta" id="openBtn">Generar Proyecto MGA</button>
<div class="overlay" id="overlay">
<div class="popup" id="popup">
<div class="popup-header">
<h3>Generador MGA</h3>
<button id="closeBtn">✕</button>
</div>
<p>Describe la idea del proyecto. La IA estructurará el MGA completo.</p>
<form method="post" action="/generar">
<textarea name="descripcion" required></textarea>
<button class="submit" type="submit">Generar ZIP MGA</button>
</form>
</div>
</div>
<script>
const hero = document.getElementById("hero");
const reveal = document.querySelector(".reveal");
const text = document.querySelector(".main-text");
hero.addEventListener("mousemove", e => {
  const b = hero.getBoundingClientRect();
  gsap.to(reveal,{ "--x": e.clientX - b.left + "px", "--y": e.clientY - b.top + "px", "--r": "220px", duration:.25 });
  text.style.opacity = .15;
});
hero.addEventListener("mouseleave", () => {
  gsap.to(reveal,{ "--r":"0px", duration:.35 });
  text.style.opacity = 1;
});
const overlay = document.getElementById("overlay");
const popup = document.getElementById("popup");
const openBtn = document.getElementById("openBtn");
const closeBtn = document.getElementById("closeBtn");
openBtn.addEventListener("click", () => {
  overlay.style.display = "flex";
  gsap.fromTo(popup,{scale:.9, opacity:0},{scale:1, opacity:1, duration:.35, ease:"power3.out"});
});
closeBtn.addEventListener("click", () => {
  gsap.to(popup,{scale:.9, opacity:0, duration:.25, onComplete:()=> overlay.style.display="none"});
});
// ⚡ cerrar al click afuera
overlay.addEventListener("click", e => { if(e.target === overlay) closeBtn.click(); });
</script>
</body>
</html>
"""


# ======================================================
# 🚀 GENERAR MGA
# ======================================================
@app.post("/generar")
def generar(descripcion: str = Form(...)):
    data = consultar_mga(descripcion)
    zip_bytes = generar_zip_completo(data)
    return StreamingResponse(
        io.BytesIO(zip_bytes),
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=Proyecto_MGA_Completo.zip"}
    )
